#needed:
#pip install python-docx
#pip install xlrd
#pip install pandas
#pip install Pyarrow
#pip install openpyxl

import docx
from docx.shared import Pt
from docx.shared import RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Cm
import pandas as pd
import re
import sys
import math
import pprint 

print("Loading Skillmatrix-Excel...")
df = pd.read_excel('Skillmatrix.xlsx', sheet_name=0)
#print(df.columns.size)
#print(len(df.index))

print("Getting target column based on name ("+sys.argv[1]+")...")
target=re.sub(r'\s', '', sys.argv[1])
pattern = re.compile(target)

targetIndex=0
for idx,col in enumerate(df.columns):
    colClean=re.sub(r'\s', '', str(col))
    #print(target)
    #print(colClean)
    if pattern.match(colClean):#colClean == target:
        print("Name found: "+re.sub(r'\s', ' ', str(col)))
        #print(idx)
        targetIndex=idx
        break
if targetIndex==0:
    print("Name not found :(")
    print("Exiting...")
    exit(0)

print("Reading Skillmatrix-Excel contents...")
allSkills={}
for index in df.index:
    #print(col)
    skill=df[df.columns[0]][index]
    value=df[df.columns[targetIndex]][index]
    #print(skill)
    #print(value)
    if str(skill) == "nan" or str(skill) == "Last Update" or str(skill) == "Legende" :
        continue
    if str(value) == "nan" and not skill.startswith(" "):
        #print("yay")
        allSkills[skill]={}
        currentMetaSkill=skill
        continue
    if str(value) == "nan" and skill[0]==" ":
        continue
    if value in (0.0,1.0,2.0,3.0,4.0):
        allSkills[currentMetaSkill][skill]=value
    else:
        print("Ignored: "+skill+" (Skill), "+value+" (Value)")
    #print(df[df.columns[0]][index])
    #print(df[df.columns[targetIndex]][index])

#pp = pprint.PrettyPrinter(indent=4)
#pp.pprint(allSkills)
#print(allSkills)


print("Loading template Word-table...")
doc = docx.Document('template.docx')
doc.tables #a list of all tables in document
#print("Retrieved value: " + doc.tables[0].cell(0, 0).text)

print("Generating Word-table...")
for metaTech, techs in allSkills.items():
    lastCell=None
    for tech,value in techs.items():
        row=doc.tables[0].add_row()
        row.height=Cm(0.75)
        c1=row.cells[0]
        if lastCell is None:
            c1.text=metaTech
            c1.paragraphs[0].paragraph_format.space_before = Pt(5)
            c1.paragraphs[0].paragraph_format.space_after = Pt(3)
            c1.paragraphs[0].paragraph_format.line_spacing = 1
            c1.paragraphs[0].runs[0].font.name='Lato'
            c1.paragraphs[0].runs[0].font.bold=True
            c1.paragraphs[0].runs[0].font.size=Pt(12)
            c1.paragraphs[0].runs[0].font.color.rgb = RGBColor(0x6B, 0x71, 0x83)
        else:
            pass
        
        c2=row.cells[1]
        c2.text=tech.strip()
        c2.paragraphs[0].paragraph_format.space_before = Pt(5)
        c2.paragraphs[0].paragraph_format.space_after = Pt(0)
        c2.paragraphs[0].paragraph_format.line_spacing = 1
        c2.paragraphs[0].runs[0].font.name='Lato'
        c2.paragraphs[0].runs[0].font.bold=False
        c2.paragraphs[0].runs[0].font.size=Pt(9)
        c2.paragraphs[0].runs[0].font.color.rgb = RGBColor(0x6B, 0x71, 0x83)
        c3=row.cells[2]
        if value == 1.0:
            c3.text="Beginner"
        elif value == 2.0:
            c3.text="Intermediate"
        elif value == 3.0:
            c3.text="Advanced"
        elif value == 4.0:
            c3.text="Expert"
        
        c3.paragraphs[0].paragraph_format.space_before = Pt(5)
        c3.paragraphs[0].paragraph_format.space_after = Pt(0)
        c3.paragraphs[0].paragraph_format.line_spacing = 1
        c3.paragraphs[0].runs[0].font.name='Lato'
        c3.paragraphs[0].runs[0].font.bold=False
        c3.paragraphs[0].runs[0].font.size=Pt(9)
        c3.paragraphs[0].runs[0].font.color.rgb = RGBColor(0x6B, 0x71, 0x83)
        if lastCell is not None:
            lastCell.merge(c1)
        lastCell=c1

print("Saving...")
doc.save("output.docx")
print("Success!")

#6B7183
